import time
import re
import os
import random
from docx import Document
from playwright.sync_api import sync_playwright

USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

def save_to_docx(question_data, filename='questions.docx'):
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    doc.add_paragraph(f"題目: {question_data.get('parsed_question','')}")
    doc.add_paragraph(f"答案: {question_data.get('stat_answer','')}")
    doc.add_paragraph(f"原始網址: {question_data.get('url','')}")
    doc.save(filename)

def parse_question(page_text):
    # 擷取「開始測驗」到「答案」之間的題目
    match = re.search(r'開始(.*?)答案', page_text, re.DOTALL)
    parsed_question = match.group(1).strip() if match else ''
    # 擷取統計答案
    stat_match = re.search(r'統計：([A-Z]\(\d+\)(?:, [A-Z]\(\d+\))*)', page_text)
    stat_answer = ''
    if stat_match:
        stat_str = stat_match.group(1)
        options = re.findall(r'([A-Z])\((\d+)\)', stat_str)
        if options:
            max_option = max(options, key=lambda x: int(x[1]))
            stat_answer = max_option[0]
    # question_data['stat_answer'] = stat_answer
    return parsed_question, stat_answer

def main(start_id=3300000, total_questions=80):
    base_url = "https://" # 網址
    print(f"準備抓取 {total_questions} 題，從ID {start_id} 開始")
    print(f"第1題: {base_url}?info=item.{start_id}")
    print(f"第{total_questions}題: {base_url}?info=item.{start_id + total_questions - 1}")
    print("-" * 60)
    docx_filename = 'question.docx'
    with sync_playwright() as p:
        # 1. 自訂 User-Agent，降低被偵測為自動化
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(user_agent=USER_AGENT)
        i = 0
        while i < total_questions:
            current_id = start_id + i
            question_url = f"{base_url}?info=item.{current_id}"
            print(f"\n正在抓取第 {i+1}/{total_questions} 題 (ID: {current_id})")
            print(f"URL: {question_url}")
            retry = 0
            while retry < 3:
                try:
                    page.goto(question_url, timeout=20000)
                    # 2. 隨機 sleep 1~2 秒，模擬真人行為
                    time.sleep(random.uniform(1, 2))
                    # 3. 模擬滾動頁面，觸發 JS 載入
                    page.evaluate("window.scrollTo(0, document.body.scrollHeight);")
                    time.sleep(0.5)
                    page.evaluate("window.scrollTo(0, 0);")
                    time.sleep(0.5)
                    # 取純文字
                    text = page.inner_text('body')
                    parsed_question, stat_answer = parse_question(text)
                    question_data = {
                        'parsed_question': parsed_question,
                        'stat_answer': stat_answer,
                        'url': question_url
                    }
                    save_to_docx(question_data, docx_filename)
                    print(f"✓ 第 {i+1} 題抓取成功")
                    break
                except Exception as e:
                    print(f"✗ 第 {i+1} 題發生錯誤: {e}")
                    retry += 1
                    if retry < 3:
                        print("等待1秒後重試...")
                        time.sleep(1)
                    else:
                        print("多次失敗，跳過本題")
            i += 1
        browser.close()

if __name__ == "__main__":
    print("Playwright 版線上測驗爬蟲")
    print("=" * 50)
    print("請先安裝 Playwright：pip install playwright")
    print("首次使用請執行：playwright install chromium")
    print("=" * 50)
    print("\n預設參數:")
    print("- 起始ID: 3300000 (第1題)")
    print("- 題目數量: 80題")
    print("- 結束ID: 3300080 (第80題)")
    user_input = input("\n是否使用預設參數？(y/n，直接按Enter使用預設): ").strip().lower()
    start_id = 3300000
    total_questions = 80
    if user_input == 'n':
        try:
            start_id = int(input("請輸入起始ID (3300000): ") or 3300000)
            total_questions = int(input("請輸入題目數量 (預設80): ") or 80)
            print(f"將抓取從ID {start_id} 開始的 {total_questions} 題")
        except ValueError:
            print("輸入格式錯誤，使用預設參數")
            start_id = 3300000
            total_questions = 80
    print(f"\n開始執行...")
    main(start_id, total_questions) 